from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Helpers ──
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),  kwargs.get('val', 'single'))
        tag.set(qn('w:sz'),   kwargs.get('sz',  '4'))
        tag.set(qn('w:space'),'0')
        tag.set(qn('w:color'),kwargs.get('color','E2E8F0'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_para(doc, text='', bold=False, size=11, color=None, space_before=0, space_after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold      = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_colored_heading(doc, text, bg_hex, text_hex='FFFFFF', size=10):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f'  {text}  ')
    run.bold           = True
    run.font.size      = Pt(size)
    run.font.color.rgb = RGBColor(*bytes.fromhex(text_hex))
    # shade the paragraph
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),  'clear')
    shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), bg_hex)
    pPr.append(shd)
    return p

def make_table(doc, headers, col_widths, rows_data, header_bg='1E40AF'):
    """
    rows_data: list of (no, rule, sub_rule, status, status_color, notes)
    """
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
        set_cell_bg(cell, header_bg)
        cell.width = col_widths[i]
        p   = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    # Data rows
    for row_data in rows_data:
        no, rule, sub, status, sc, notes = row_data
        row_cells = table.add_row().cells

        # #
        row_cells[0].width = col_widths[0]
        p = row_cells[0].paragraphs[0]
        r = p.add_run(no)
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Rule
        row_cells[1].width = col_widths[1]
        p = row_cells[1].paragraphs[0]
        r = p.add_run(rule)
        r.bold = True; r.font.size = Pt(9)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(1)
        if sub:
            p2 = row_cells[1].add_paragraph(sub)
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(3)
            for run in p2.runs:
                run.font.size      = Pt(8)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # Status
        row_cells[2].width = col_widths[2]
        p = row_cells[2].paragraphs[0]
        r = p.add_run(status)
        r.bold = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(*bytes.fromhex(sc))
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Notes (empty fill-in area)
        row_cells[3].width = col_widths[3]
        p = row_cells[3].paragraphs[0]
        r = p.add_run(notes)
        r.font.size      = Pt(8.5)
        r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        set_cell_bg(row_cells[3], 'FAFAFA')

    return table


# ════════════════════════════════════════
#  TITLE
# ════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run('排班 Constraints 討論清單')
r.bold = True; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

# ── Meeting info table ──
info = doc.add_table(rows=1, cols=3)
info.style = 'Table Grid'
labels = ['日期：', '出席：', '紀錄：']
for cell, lbl in zip(info.rows[0].cells, labels):
    p = cell.paragraphs[0]
    r = p.add_run(lbl)
    r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    p.add_run('_' * 18).font.size = Pt(9)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    set_cell_bg(cell, 'F1F5F9')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Summary ──
add_para(doc, '本次討論範圍', bold=True, size=9, color='64748B', space_before=0, space_after=2)
sum_table = doc.add_table(rows=1, cols=3)
sum_table.style = 'Table Grid'
sum_data = [
    ('9',  'Hard Constraints\n（絕對不能違反）', 'DC2626'),
    ('4',  'Soft Constraints\n（盡量符合）',     'D97706'),
    ('2',  '特別邏輯\n（CC Combine）',           '1E40AF'),
]
for cell, (num, lbl, color) in zip(sum_table.rows[0].cells, sum_data):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(num + '\n')
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    r2 = p.add_run(lbl)
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ════════════════════════════════════════
#  HARD CONSTRAINTS
# ════════════════════════════════════════
add_colored_heading(doc, '🔴  HARD CONSTRAINTS — 絕對不能違反，違反即排班失敗', 'FEE2E2', 'DC2626')

col_w = [Cm(1.2), Cm(7.8), Cm(2.2), Cm(5.4)]
hard_rows = [
    ('H1', '老師不能同一時段教兩班',
     '同一老師同一天同一時段只能有一個班（主教 Lec1）',
     '✓ 確認', '15803D', ''),
    ('H2', '課室不能同時用兩班',
     '同一課室同一時段只排一班；不可超出座位上限',
     '✓ 確認', '15803D', ''),
    ('H3', 'Core 科上課地點 = 學生報名 Campus',
     '學生在哪個 center 報名，Core 科就在那個 center 上課',
     '✓ 確認', '15803D', ''),
    ('H4', '老師每週 Loading 上限 6 班',
     '1 loading = 1 班（唯論 2hr 定 4hr）；個別老師可另行設定',
     '✓ 確認', '15803D', ''),
    ('H5', 'CC Combine 條件：同科目 + 同語言',
     '只有同科目代號且同授課語言（中文/英文）才能合班',
     '✓ 確認', '15803D', ''),
    ('H6', 'TKO Center 特別時段',
     'TKO 只可排 10:00–14:00 或 15:00–19:00（冷氣費問題）；老師+學生均適用',
     '✓ 確認', '15803D', ''),
    ('H7', '老師需具備教授該科目資格',
     'Teacher Load Table 內有登記，且未超出個人 Loading 上限',
     '✓ 確認', '15803D', ''),
    ('H8', '老師同一天不能去兩個 Center',
     '絕對不允許，不設例外；分配老師時必須逐一檢查',
     '✓ 確認', '15803D', ''),
    ('H9', '學生同一天不能去兩個 Center',
     '排班時需確保：同一學生當天所有課都在同一個 Center',
     '✓ 確認', '15803D', ''),
]
make_table(doc, ['#', '規則', '狀態', '問題 / 備注（會議填寫）'], col_w, hard_rows, header_bg='DC2626')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════
#  SOFT CONSTRAINTS
# ════════════════════════════════════════
add_colored_heading(doc, '🟡  SOFT CONSTRAINTS — 盡量符合，違反時發出警告', 'FEF3C7', 'B45309')

soft_rows = [
    ('S1', '老師不可用時間（Preference）',
     '老師申報不許可的時段，系統優先跳過；萬不得已才排入並發警告',
     '？ 待確認', 'B45309', '老師如何申報不可用時間？格式待定'),
    ('S2', '不要天地堂',
     '同半天（AM 或 PM）內，兩堂課之間空堂 >2小時 = 天地堂；午飯不計；老師+學生均適用',
     '✓ 確認', '15803D', ''),
    ('S3', '班別平均分佈至每天（輪流發牌）',
     '同科目多班，輪流從不同日子（Mon→Tue→Thu→Wed→Fri）開始排，避免集中某天',
     '✓ 確認', '15803D', ''),
    ('S4', 'CC Combine 地點選人數最多的 Center',
     '減少學生移動；平局時選距 SSP/CSW 較近者；受 H9 約束',
     '✓ 確認', '15803D', ''),
]
make_table(doc, ['#', '規則', '狀態', '問題 / 備注（會議填寫）'], col_w, soft_rows, header_bg='D97706')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════
#  CC COMBINE LOGIC
# ════════════════════════════════════════
add_colored_heading(doc, '🔵  CC COMBINE 特別邏輯', 'DBEAFE', '1E40AF')

logic_rows = [
    ('L1', '方法 A：先排 Core，再為 CC 找合適日子（優先）',
     'CC combine 班只排在「所有相關學生當天沒有 Core 科在其他 Center」的日子；找到即完成',
     '✓ 確認', '15803D', ''),
    ('L2', '方法 B：調整 Combine 地點（備用）',
     '方法 A 無解時，改以次多學生 Center 為 combine 地點重新搜尋；全部失敗 → ERROR，人手處理',
     '✓ 確認', '15803D', ''),
]
make_table(doc, ['#', '邏輯', '狀態', '問題 / 備注（會議填寫）'], col_w, logic_rows, header_bg='1E40AF')

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ════════════════════════════════════════
#  OPEN ISSUES
# ════════════════════════════════════════
add_colored_heading(doc, '📋  會議提出的問題 / 跟進事項', 'F5F3FF', '6D28D9')

issue_table = doc.add_table(rows=6, cols=4)
issue_table.style = 'Table Grid'

# Header
hdrs = ['#', '相關規則', '問題描述', '負責人 / 期限']
issue_widths = [Cm(0.9), Cm(2.2), Cm(9.5), Cm(4.0)]
for i, (cell, hdr) in enumerate(zip(issue_table.rows[0].cells, hdrs)):
    set_cell_bg(cell, '6D28D9')
    cell.width = issue_widths[i]
    p = cell.paragraphs[0]
    r = p.add_run(hdr)
    r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# Pre-fill row 1 with known issue
pre = [('1', 'S1', '老師不可用時間的輸入格式待定（填表格？系統選擇？粒度？）', '')]
for row_idx, (no, rule, desc, owner) in enumerate(pre, start=1):
    cells = issue_table.rows[row_idx].cells
    for i, (cell, val, w) in enumerate(zip(cells, [no, rule, desc, owner], issue_widths)):
        cell.width = w
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)
        if i == 0:
            r.bold = True
            r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(14)  # leave space to write

# Remaining blank rows
for row_idx in range(2, 6):
    cells = issue_table.rows[row_idx].cells
    for i, (cell, w) in enumerate(zip(cells, issue_widths)):
        cell.width = w
        p = cell.paragraphs[0]
        if i == 0:
            r = p.add_run(str(row_idx))
            r.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(14)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ════════════════════════════════════════
#  SIGN OFF
# ════════════════════════════════════════
sign_table = doc.add_table(rows=1, cols=2)
sign_table.style = 'Table Grid'
sign_data = ['出席者確認（簽名 / 日期）', '下次會議日期：______ 年 ______ 月 ______ 日']
for cell, lbl in zip(sign_table.rows[0].cells, sign_data):
    p = cell.paragraphs[0]
    r = p.add_run(lbl)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(28)
    set_cell_bg(cell, 'F8FAFC')

# Footer note
add_para(doc, 'HKIT 排班系統 · V4 · 2026年5月', size=8, color='94A3B8',
         space_before=10, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
out = r'C:\Users\kmksy\Desktop\HKIT\Demo_sharing\09_timetable\scheduling-meeting.docx'
doc.save(out)
print(f'Saved: {out}')
